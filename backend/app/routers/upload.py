"""
Document Upload Router
Accepts PDF, DOCX, TXT, and CSV file uploads for competitive intelligence analysis.
Extracts text content and ingests it as a snapshot for the specified competitor.
"""
import uuid
import io
from datetime import datetime, timezone
from typing import Annotated, Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from sqlalchemy.orm import Session

from app.database import get_db
from app.dependencies.auth import get_current_user
from app.models.user import User
from app.models.competitor import Competitor
from app.models.snapshot import Snapshot, SourceType
from app.services.vector_store import vector_store

router = APIRouter(prefix="/upload", tags=["upload"])

# Maximum file size: 20MB
MAX_FILE_SIZE = 20 * 1024 * 1024

# Supported MIME types and extensions
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".md", ".json"}
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/csv",
    "text/markdown",
    "application/json",
    "application/octet-stream",  # Fallback for untyped uploads
}


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF bytes using fpdf2's built-in reader or pdfplumber fallback."""
    text = ""
    try:
        # Try PyPDF2 / pypdf first (most common)
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text() or ""
                text += page_text + "\n"
        except ImportError:
            pass

        # Fallback: try pdfplumber
        if not text.strip():
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n"
            except ImportError:
                pass

        # Fallback: try fitz (PyMuPDF)
        if not text.strip():
            try:
                import fitz
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except ImportError:
                pass

        if not text.strip():
            raise ValueError("No PDF reader library available. Install pypdf, pdfplumber, or PyMuPDF.")

    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {e}")

    return text.strip()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except ImportError:
        raise ValueError("python-docx library not installed. Install it with: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def _extract_text_from_csv(file_bytes: bytes) -> str:
    """Extracts readable text from CSV bytes, converting rows to a readable format."""
    import csv
    try:
        text_content = file_bytes.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text_content))
        rows = list(reader)
        if not rows:
            return ""

        # Convert to markdown table format for better LLM processing
        lines = []
        headers = rows[0] if rows else []
        if headers:
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows[1:100]:  # Cap at 100 rows
            padded = row + [""] * (len(headers) - len(row))
            lines.append("| " + " | ".join(padded[:len(headers)]) + " |")

        return "\n".join(lines)
    except Exception as e:
        raise ValueError(f"Failed to parse CSV: {e}")


def _extract_text_from_upload(file_bytes: bytes, filename: str, content_type: str) -> str:
    """Routes to the correct text extractor based on file type."""
    ext = ""
    if "." in filename:
        ext = "." + filename.rsplit(".", 1)[-1].lower()

    if ext == ".pdf" or "pdf" in content_type:
        return _extract_text_from_pdf(file_bytes)
    elif ext == ".docx" or "wordprocessingml" in content_type:
        return _extract_text_from_docx(file_bytes)
    elif ext == ".csv" or "csv" in content_type:
        return _extract_text_from_csv(file_bytes)
    elif ext in (".txt", ".md", ".json") or "text/" in content_type or "json" in content_type:
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        # Best effort: try as plain text
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
            if text.strip():
                return text
        except Exception:
            pass
        raise ValueError(f"Unsupported file type: {ext or content_type}")


@router.post("/document")
async def upload_document(
    file: UploadFile = File(...),
    competitor_id: Optional[str] = Form(None),
    source_type: Optional[str] = Form("REVIEW"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Uploads a document (PDF, DOCX, TXT, CSV) and extracts text content.
    If competitor_id is provided, ingests the content as a snapshot for that competitor.
    Otherwise, returns the extracted text for the user to review.
    """
    filename = file.filename or "upload.txt"

    # Read file content with size limit
    file_bytes = await file.read()
    if len(file_bytes) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large ({len(file_bytes) / 1024 / 1024:.1f}MB). Maximum: {MAX_FILE_SIZE / 1024 / 1024:.0f}MB",
        )

    if len(file_bytes) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file is empty.",
        )

    # Universal Document Text Extraction
    try:
        from app.services.document_parser import extract_text_from_any_document
        extracted_text = extract_text_from_any_document(filename, file_bytes, file.content_type)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to extract document text: {e}",
        )

    if not extracted_text or len(extracted_text.strip()) < 50:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Could not extract sufficient text from the uploaded file. The document may be image-based or empty.",
        )

    result = {
        "filename": filename,
        "file_size_bytes": len(file_bytes),
        "extracted_text_length": len(extracted_text),
        "extracted_text_preview": extracted_text[:500] + ("..." if len(extracted_text) > 500 else ""),
    }

    # If competitor_id is provided, ingest as a snapshot
    if competitor_id:
        try:
            comp_uuid = uuid.UUID(competitor_id)
        except ValueError:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid competitor_id format.",
            )

        competitor = db.get(Competitor, comp_uuid)
        if not competitor or competitor.user_id != current_user.id:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Competitor not found.",
            )

        # Determine source type
        st_map = {"PRICING": SourceType.PRICING, "REVIEW": SourceType.REVIEW, "NEWS": SourceType.NEWS}
        snap_source = st_map.get((source_type or "REVIEW").upper(), SourceType.REVIEW)

        # Create snapshot from uploaded document
        import hashlib
        content_hash = hashlib.sha256(extracted_text.encode("utf-8")).hexdigest()

        from app.services.scraper import sanitize_text_content
        extracted_text_safe = sanitize_text_content(extracted_text)

        snapshot = Snapshot(
            competitor_id=competitor.id,
            source_type=snap_source,
            raw_content=extracted_text_safe,
            content_hash=content_hash,
            is_stale=False,
            fetched_at=datetime.now(timezone.utc),
        )
        db.add(snapshot)
        db.commit()
        db.refresh(snapshot)

        # Index in FAISS vector store
        chunks_added = 0
        try:
            chunks_added = vector_store.add_snapshot_chunks(
                snapshot_id=str(snapshot.id),
                competitor_id=str(competitor.id),
                source_type=snap_source.value,
                fetched_at=snapshot.fetched_at.isoformat(),
                text=extracted_text,
            )
        except Exception as vec_err:
            print(f"[Upload] FAISS indexing warning: {vec_err}", flush=True)

        result.update({
            "status": "ingested",
            "snapshot_id": str(snapshot.id),
            "competitor_id": str(competitor.id),
            "competitor_name": competitor.name,
            "source_type": snap_source.value,
            "chunks_indexed": chunks_added,
        })
    else:
        result["status"] = "extracted"
        result["extracted_text"] = extracted_text[:10000]  # Cap preview for response size
        result["message"] = "Text extracted successfully. Provide competitor_id to ingest as a snapshot."

    return result
